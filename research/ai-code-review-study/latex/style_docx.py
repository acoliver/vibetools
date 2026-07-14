#!/usr/bin/env python3
"""Apply restrained technical-report styling to a Pandoc-generated DOCX."""

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x17, 0x20, 0x33)
MUTED = RGBColor(0x5D, 0x68, 0x78)
ACCENT = RGBColor(0x24, 0x57, 0xA7)
ACCENT_LIGHT = "EAF1FB"
TABLE_ALT = "F5F7FA"
RULE = "CBD5E1"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def set_repeat_header(row):
    tr_props = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_props.append(marker)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Charter"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 14),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 20, INK, 18, 8),
        ("Heading 2", 14, ACCENT, 14, 5),
        ("Heading 3", 11, INK, 10, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1", "Heading 2"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Caption", "Intense Quote"):
        if name in doc.styles:
            style = doc.styles[name]
            style.font.name = "Aptos"
            style.font.size = Pt(9)
            style.font.color.rgb = MUTED

    if "Executive Callout" not in doc.styles:
        style = doc.styles.add_style("Executive Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout = doc.styles["Executive Callout"]
    callout.font.name = "Aptos"
    callout.font.size = Pt(10)
    callout.font.color.rgb = INK
    callout.paragraph_format.left_indent = Inches(0.22)
    callout.paragraph_format.right_indent = Inches(0.22)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(7)


def configure_sections(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.32)
        header = section.header.paragraphs[0]
        header.text = "AI Code Review Systems  |  Technical Research Report"
        header.style = doc.styles["Normal"]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)
            run.font.color.rgb = MUTED
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(footer)


def style_tables(doc):
    for table in doc.tables:
        table.autofit = True
        if table.rows:
            set_repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                shade(cell, ACCENT_LIGHT if row_index == 0 else (TABLE_ALT if row_index % 2 == 0 else "FFFFFF"))
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(1.5)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = "Aptos"
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = INK
                        if row_index == 0:
                            run.font.bold = True
                            run.font.color.rgb = ACCENT


def style_paragraphs(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name.startswith("Heading"):
            set_keep_with_next(paragraph)
        if paragraph.style.name in {"Block Text", "Intense Quote"} or text.startswith("This brief stands alone."):
            paragraph.style = doc.styles["Executive Callout"]
            props = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), ACCENT_LIGHT)
            props.append(shading)
            border = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:color"), "2457A7")
            left.set(qn("w:space"), "6")
            border.append(left)
            props.append(border)
        if text.startswith("Figure ") or text.startswith("*Figure "):
            paragraph.style = doc.styles["Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def constrain_images(doc):
    for shape in doc.inline_shapes:
        max_width = Inches(6.65)
        max_height = Inches(7.0)
        ratio = min(max_width / shape.width, max_height / shape.height, 1)
        shape.width = int(shape.width * ratio)
        shape.height = int(shape.height * ratio)


def main():
    import sys
    source = Path(sys.argv[1])
    destination = Path(sys.argv[2])
    doc = Document(source)
    configure_styles(doc)
    configure_sections(doc)
    style_tables(doc)
    style_paragraphs(doc)
    constrain_images(doc)
    core = doc.core_properties
    core.title = "AI Code Review Systems at VybeStack"
    core.subject = "Comparative quality, reliability, provider behavior, and duplicate mitigation"
    core.author = "LLxprt Research"
    doc.save(destination)
    print(f"styled {destination}")


if __name__ == "__main__":
    main()
